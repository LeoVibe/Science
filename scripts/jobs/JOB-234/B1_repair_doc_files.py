#!/usr/bin/env python3
"""JOB-234 Phase 1：DOC 格式重抽。
讀 _repair_manifest.json，對 source_type=doc_format 的項目用 python-docx 重抽，
覆寫整合版 MD（保留 frontmatter，替換內容段落，更新 quality_flags）。
"""
import os, json, re
from datetime import datetime, timezone, timedelta

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
os.chdir(ROOT)

MANIFEST_PATH = "scripts/jobs/JOB-234/_repair_manifest.json"
REPORT_PATH   = "scripts/jobs/JOB-234/_doc_repair_log.json"


def _extract_jpegs_from_ole(stream_data):
    """從 OLE WordDocument stream 提取 JPEG bytes（> 5KB 才回傳）。"""
    images = []
    i = 0
    while i < len(stream_data) - 3:
        if stream_data[i:i+3] == b'\xff\xd8\xff':
            end = stream_data.find(b'\xff\xd9', i + 3)
            if end != -1:
                chunk = stream_data[i:end + 2]
                if len(chunk) > 5000:
                    images.append(chunk)
                i = end + 2
                continue
        i += 1
    return images


def _ocrmac_images(image_bytes_list):
    """用 macOS Vision (ocrmac) OCR 一組圖片 bytes，回傳合併文字。"""
    import tempfile, shutil
    try:
        from ocrmac import ocrmac as ocrmac_lib
    except ImportError:
        raise RuntimeError("ocrmac 未安裝：pip3.11 install ocrmac")

    tmpdir = tempfile.mkdtemp()
    parts = []
    try:
        for idx, img_bytes in enumerate(image_bytes_list):
            img_path = os.path.join(tmpdir, f'img_{idx}.jpg')
            with open(img_path, 'wb') as f:
                f.write(img_bytes)
            ann = ocrmac_lib.OCR(
                img_path,
                language_preference=["zh-Hant", "zh-Hans", "en-US"],
                recognition_level="accurate",
            ).recognize()
            text = '\n'.join(a[0] for a in ann if a[0].strip())
            if text.strip():
                parts.append(f'## 圖片 {idx + 1}\n\n{text}')
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)
    return '\n\n---\n\n'.join(parts)


def extract_doc_text(doc_path):
    """抽取 DOC/DOCX 文字。
    .docx → python-docx
    .doc  → ① textutil（有文字層）→ ② OLE JPEG + ocrmac（嵌圖型 fallback）
    """
    import subprocess, tempfile

    if doc_path.lower().endswith('.doc'):
        # ① 先試 textutil（有文字層的 .doc 直接成功）
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
            tmp_path = tmp.name
        try:
            result = subprocess.run(
                ['textutil', '-convert', 'txt', '-output', tmp_path, doc_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0:
                with open(tmp_path, encoding='utf-8', errors='replace') as f:
                    text = f.read().strip()
                if text:
                    return text
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

        # ② textutil 空白 → 圖片嵌入型 .doc：OLE 提取 JPEG + macOS Vision OCR
        try:
            import olefile
        except ImportError:
            raise RuntimeError("olefile 未安裝：pip3.11 install olefile")
        ole = olefile.OleFileIO(doc_path)
        stream_data = ole.openstream('WordDocument').read()
        ole.close()
        images = _extract_jpegs_from_ole(stream_data)
        if not images:
            raise RuntimeError("OLE 中未找到可用圖片（JPEG > 5KB）")
        text = _ocrmac_images(images)
        if not text.strip():
            raise RuntimeError("ocrmac OCR 輸出空白")
        return text
    else:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx 未安裝。請執行：pip3.11 install python-docx")
        doc = Document(doc_path)

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(' | '.join(cells))
            if rows:
                paragraphs.append('\n'.join(rows))

        return '\n\n'.join(paragraphs)


def update_md_file(md_path, new_content, source_type):
    """保留 frontmatter，替換內文，更新 quality_flags 和 char_count。"""
    with open(md_path, encoding='utf-8') as f:
        original = f.read()

    # 分割 frontmatter 與 content
    m = re.match(r'^(---\n.*?\n---\n)(.*)', original, re.S)
    if not m:
        raise ValueError(f"無法解析 frontmatter: {md_path}")

    frontmatter_block = m.group(1)

    # 更新 quality_flags：移除 extract_failed / paper_empty，加入 docx_extracted
    def update_flags(fm_block):
        flags_to_remove = {'extract_failed', 'paper_empty'}
        flags_to_add    = {'docx_extracted'}
        lines = fm_block.split('\n')
        new_lines = []
        in_flags = False
        added = False
        for line in lines:
            if 'quality_flags:' in line:
                in_flags = True
                new_lines.append(line)
                continue
            if in_flags:
                if line.strip().startswith('- '):
                    flag = line.strip()[2:].strip()
                    if flag in flags_to_remove:
                        continue  # 移除
                    if not added:
                        for f in sorted(flags_to_add):
                            new_lines.append(f'  - {f}')
                        added = True
                    new_lines.append(line)
                else:
                    if not added:
                        for f in sorted(flags_to_add):
                            new_lines.append(f'  - {f}')
                        added = True
                    in_flags = False
                    new_lines.append(line)
            else:
                new_lines.append(line)
        return '\n'.join(new_lines)

    frontmatter_block = update_flags(frontmatter_block)

    # 更新 char_count
    char_count = len(new_content)
    frontmatter_block = re.sub(
        r'^char_count:\s*\d+',
        f'char_count: {char_count}',
        frontmatter_block,
        flags=re.M
    )

    # 更新 integration.method
    frontmatter_block = re.sub(
        r'(  method:\s*).*',
        r'\1"JOB-234 python-docx 重抽"',
        frontmatter_block
    )

    now = datetime.now(timezone(timedelta(hours=8))).strftime('%Y-%m-%d')
    frontmatter_block = re.sub(
        r'(  integrated_date:\s*).*',
        rf'\g<1>{now}',
        frontmatter_block
    )

    new_md = frontmatter_block + '\n' + new_content

    # 原子寫入
    tmp_path = md_path + '.tmp'
    with open(tmp_path, 'w', encoding='utf-8') as f:
        f.write(new_md)
    os.replace(tmp_path, md_path)
    return char_count


def main():
    with open(MANIFEST_PATH, encoding='utf-8') as f:
        manifest = json.load(f)

    items = [it for it in manifest['items'] if it['source_type'] == 'doc_format']
    print(f'Phase 1：處理 {len(items)} 份 DOC 格式')

    log = []
    success = failed = 0

    for item in items:
        exam_id  = item['exam_id']
        md_path  = item['md_path']
        orig_files = item.get('original_files', [])

        # 找試卷（優先 試卷，其次全部）
        exam_files = [f for f in orig_files if '試卷' in os.path.basename(f)]
        ans_files  = [f for f in orig_files if '答案' in os.path.basename(f)]
        target_files = exam_files + ans_files if exam_files else orig_files

        if not target_files:
            print(f'  ⚠ {exam_id}: 無原始檔，跳過')
            log.append({'exam_id': exam_id, 'status': 'skipped', 'reason': 'no_original_files'})
            continue

        extracted_parts = []
        for doc_path in target_files:
            try:
                text = extract_doc_text(doc_path)
                if text.strip():
                    extracted_parts.append(f'## {os.path.basename(doc_path)}\n\n{text}')
            except Exception as e:
                print(f'  ❌ {exam_id} [{os.path.basename(doc_path)}]: {e}')

        if not extracted_parts:
            print(f'  ❌ {exam_id}: 所有原始檔抽取失敗')
            log.append({'exam_id': exam_id, 'status': 'failed', 'reason': 'extraction_error'})
            failed += 1
            continue

        new_content = '\n\n---\n\n'.join(extracted_parts)
        try:
            char_count = update_md_file(md_path, new_content, 'doc_format')
            print(f'  ✅ {exam_id}: char_count={char_count}')
            log.append({'exam_id': exam_id, 'status': 'success', 'char_count': char_count})
            # 更新 manifest 狀態
            item['repair_status'] = 'success'
            success += 1
        except Exception as e:
            print(f'  ❌ {exam_id}: MD 更新失敗 {e}')
            log.append({'exam_id': exam_id, 'status': 'failed', 'reason': str(e)})
            failed += 1

    print(f'\n=== Phase 1 完成：成功 {success} / 失敗 {failed} ===')

    # 更新 manifest
    with open(MANIFEST_PATH, 'w', encoding='utf-8') as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)

    with open(REPORT_PATH, 'w', encoding='utf-8') as f:
        json.dump({'success': success, 'failed': failed, 'items': log},
                  f, ensure_ascii=False, indent=2)
    print(f'Log 寫入：{REPORT_PATH}')


if __name__ == '__main__':
    main()
